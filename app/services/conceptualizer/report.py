"""DOCX report generator for Conceptualizator three-layer output."""
import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import ConceptualizationOutput


def generate_concept_docx(
    output: ConceptualizationOutput,
    meta: dict | None = None,
) -> io.BytesIO:
    """Build a DOCX report from the three-layer conceptualization output.

    Returns an io.BytesIO positioned at start, ready for upload.
    """
    doc = Document()

    # ── Helpers ──────────────────────────────────────────────────────────────

    def _heading(text: str, level: int = 2) -> None:
        doc.add_heading(text, level=level)

    def _para(text: str) -> None:
        doc.add_paragraph(text)

    def _bold_para(text: str) -> None:
        p = doc.add_paragraph()
        p.add_run(text).bold = True

    def _label(label: str, value: str) -> None:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    # ── Title ─────────────────────────────────────────────────────────────────

    title = doc.add_heading("Концептуализация случая", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_str = (meta or {}).get("date", datetime.now(timezone.utc).strftime("%Y-%m-%d"))
    _para(f"Дата: {date_str}")
    session_short = output.session_id[:8] if output.session_id else "—"
    _para(f"Сессия: {session_short}")
    doc.add_paragraph()

    # ── Layer A ───────────────────────────────────────────────────────────────

    _heading("Layer A — Концептуальная модель")

    _label("Ведущая гипотеза", output.layer_a.leading_formulation)
    _label("Доминирующий слой", output.layer_a.dominant_layer.value)

    doc.add_paragraph()
    _bold_para("Конфигурация системы (петли обратной связи):")
    _para(output.layer_a.configuration_summary)

    doc.add_paragraph()
    _bold_para("Цена системы:")
    _para(output.layer_a.system_cost)

    if output.layer_a.supporting_points:
        doc.add_paragraph()
        _bold_para("Опорные гипотезы:")
        for point in output.layer_a.supporting_points:
            doc.add_paragraph(point, style="List Bullet")

    # ── Layer B ───────────────────────────────────────────────────────────────

    doc.add_paragraph()
    _heading("Layer B — Мишени вмешательства")

    for i, target in enumerate(output.layer_b.targets, 1):
        _bold_para(f"{i}. {target.layer}  (приоритет {target.priority})")
        _label("Направление", target.direction)
        _label("Обоснование", target.rationale)
        doc.add_paragraph()

    if output.layer_b.sequencing_notes:
        _bold_para("Рекомендуемая последовательность:")
        _para(output.layer_b.sequencing_notes)

    # ── Layer C ───────────────────────────────────────────────────────────────

    doc.add_paragraph()
    _heading("Layer C — Нарратив для клиента")

    _label("Метафора", f"«{output.layer_c.core_metaphor}»")
    doc.add_paragraph()
    _para(output.layer_c.narrative)
    doc.add_paragraph()
    _bold_para("Направление изменения:")
    _para(output.layer_c.direction_of_change)

    # ── Save ──────────────────────────────────────────────────────────────────

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
