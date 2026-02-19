"""Screen v2 report generation.

Three public callables:
  generate_full_report  — async; calls Claude twice, assembles report_json + report_text
  format_report_txt     — pure; formats report_json into a readable Russian plain-text
  generate_report_docx  — async; builds a professional DOCX from report_json
"""
from __future__ import annotations

import io
import json
import logging
from datetime import datetime, timezone

logger = logging.getLogger(__name__)

_SONNET = "claude-sonnet-4-5-20250929"

# ---------------------------------------------------------------------------
# Axis / Layer metadata used across all formatters
# ---------------------------------------------------------------------------

_AXIS_META = {
    "A1": {
        "name": "Активация",
        "negative": "сниженная",
        "neutral": "стабильная",
        "positive": "повышенная",
    },
    "A2": {
        "name": "Неопределённость",
        "negative": "избегание",
        "neutral": "нейтральная",
        "positive": "исследование",
    },
    "A3": {
        "name": "Импульс",
        "negative": "импульсивность",
        "neutral": "гибкая пауза",
        "positive": "избыточное удержание",
    },
    "A4": {
        "name": "Временная ориентация",
        "negative": "короткий цикл",
        "neutral": "смешанная",
        "positive": "горизонт",
    },
}

_LAYER_META = {
    "L0": "Энергетический",
    "L1": "Поведенческий",
    "L2": "Аффективный",
    "L3": "Социальный",
    "L4": "Когнитивный",
}


def _axis_label(axis_key: str, value: float) -> str:
    meta = _AXIS_META.get(axis_key, {})
    if value < -0.3:
        pole = meta.get("negative", "−")
    elif value > 0.3:
        pole = meta.get("positive", "+")
    else:
        pole = meta.get("neutral", "нейтральная")
    return pole


def _rigidity_label(total: float) -> str:
    if total < 0.3:
        return "low"
    if total < 0.6:
        return "medium"
    return "high"


# ---------------------------------------------------------------------------
# 1. generate_full_report
# ---------------------------------------------------------------------------

async def generate_full_report(state: dict, claude_client) -> dict:
    """Build the complete report for a finished screening assessment.

    Calls Claude twice (report generator + session bridge), assembles
    report_json, and produces report_text via format_report_txt().

    Args:
        state:         Current assessment state dict from the orchestrator.
        claude_client: anthropic.AsyncAnthropic instance.

    Returns:
        {"report_json": dict, "report_text": str}
    """
    from app.services.screen.prompts import (
        REPORT_GENERATOR_PROMPT,
        SESSION_BRIDGE_PROMPT,
        assemble_prompt,
    )

    # ---- 1. Structural report (Claude sonnet) ----------------------------
    report_context = {
        "AxisVector": state.get("axis_vector", {}),
        "LayerVector": state.get("layer_vector", {}),
        "TensionMatrix": state.get("tension_matrix", {}),
        "RigidityIndex": state.get("rigidity", {}),
        "DominantCells": state.get("dominant_cells", []),
        "Confidence": state.get("confidence", 0.0),
    }
    report_user = assemble_prompt("report", report_context)
    structural_report = await _call_claude(
        claude_client,
        system=REPORT_GENERATOR_PROMPT,
        user_content=report_user,
        model=_SONNET,
        max_tokens=2000,
    )
    if not structural_report:
        structural_report = "(Отчёт недоступен — ошибка генерации)"

    # ---- 2. Session bridge / interview protocol (Claude sonnet) ----------
    bridge_context = {
        "AxisVector": state.get("axis_vector", {}),
        "LayerVector": state.get("layer_vector", {}),
        "DominantCells": state.get("dominant_cells", []),
        "RigidityIndex": state.get("rigidity", {}),
        "Confidence": state.get("confidence", 0.0),
    }
    bridge_user = assemble_prompt("session_bridge", bridge_context)
    bridge_raw = await _call_claude(
        claude_client,
        system=SESSION_BRIDGE_PROMPT,
        user_content=bridge_user,
        model=_SONNET,
        max_tokens=1500,
    )
    interview_protocol: dict = {}
    if bridge_raw:
        try:
            interview_protocol = _parse_json(bridge_raw)
        except Exception:
            logger.warning("[screen] Session bridge JSON parse failed; leaving empty")
            interview_protocol = {"axis_verification": [], "layer_exploration": [], "functional_context": []}

    # ---- 3. Assemble report_json -----------------------------------------
    report_json = {
        "assessment_id": state.get("assessment_id", ""),
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "axis_vector": state.get("axis_vector", {}),
        "layer_vector": state.get("layer_vector", {}),
        "tension_matrix": state.get("tension_matrix", {}),
        "dominant_cells": state.get("dominant_cells", []),
        "rigidity": state.get("rigidity", {}),
        "confidence": state.get("confidence", 0.0),
        "phases": {
            "phase1_questions": 6,
            "phase2_questions": state.get("phase2_questions", 0),
            "phase3_questions": state.get("phase3_questions", 0),
        },
        "structural_report": structural_report,
        "interview_protocol": interview_protocol,
    }

    report_text = format_report_txt(report_json)
    return {"report_json": report_json, "report_text": report_text}


# ---------------------------------------------------------------------------
# 2. format_report_txt
# ---------------------------------------------------------------------------

def format_report_txt(report_json: dict) -> str:
    """Render report_json as a readable Russian plain-text string."""
    axis_v = report_json.get("axis_vector", {})
    layer_v = report_json.get("layer_vector", {})
    dominant = report_json.get("dominant_cells", [])
    rigidity = report_json.get("rigidity", {})
    tension = report_json.get("tension_matrix", {})
    confidence = report_json.get("confidence", 0.0)
    structural = report_json.get("structural_report", "")
    protocol = report_json.get("interview_protocol", {})
    ts = report_json.get("timestamp", "")
    try:
        date_str = datetime.fromisoformat(ts).strftime("%d.%m.%Y")
    except Exception:
        date_str = ts[:10] if ts else "—"

    sep = "═" * 47
    lines: list[str] = [
        sep,
        "PsycheOS Screening v2 — Структурный профиль",
        f"Дата: {date_str}",
        sep,
        "",
        "▸ ПРОФИЛЬ ОСЕЙ РЕГУЛЯЦИИ",
    ]

    for key in ["A1", "A2", "A3", "A4"]:
        val = axis_v.get(key, 0.0)
        meta = _AXIS_META[key]
        label = _axis_label(key, val)
        lines.append(f"  {meta['name']} ({key}): {val:+.3f} — {label}")

    lines += ["", "▸ ДОМИНИРУЮЩИЕ СЛОИ"]
    # Sort layers by |value| descending
    sorted_layers = sorted(layer_v.items(), key=lambda kv: abs(kv[1]), reverse=True)
    for rank, (lkey, lval) in enumerate(sorted_layers, start=1):
        layer_name = _LAYER_META.get(lkey, lkey)
        lines.append(f"  {rank}. {layer_name} ({lkey}): {lval:+.3f}")

    lines += ["", "▸ КЛЮЧЕВЫЕ СОЧЕТАНИЯ"]
    for rank, cell in enumerate(dominant[:3], start=1):
        cell_val = tension.get(cell, 0.0)
        lines.append(f"  {rank}. {cell}: {cell_val:+.3f}")

    rig_total = rigidity.get("total", 0.0)
    rig_label = _rigidity_label(rig_total)
    rig_pol = rigidity.get("polarization", 0.0)
    rig_var = rigidity.get("low_variance", 0.0)
    rig_rep = rigidity.get("strategy_repetition", 0.0)
    lines += [
        "",
        "▸ ИНДЕКС ГИБКОСТИ",
        f"  Ригидность: {rig_label} ({rig_total:.2f})",
        f"  Поляризация: {rig_pol:.2f} | Стабильность: {rig_var:.2f} | Повторяемость: {rig_rep:.2f}",
        "",
        f"▸ УВЕРЕННОСТЬ: {confidence * 100:.0f}%",
        "",
        "▸ ПОЯСНЕНИЕ",
    ]
    lines.append(structural)

    lines += ["", "▸ ОРИЕНТИРЫ ДЛЯ ПЕРВОЙ СЕССИИ"]
    if isinstance(protocol, dict):
        for section_key, section_label in [
            ("axis_verification", "Верификация осей"),
            ("layer_exploration", "Исследование слоёв"),
            ("functional_context", "Функциональный контекст"),
        ]:
            questions = protocol.get(section_key, [])
            if questions:
                lines.append(f"\n  {section_label}:")
                for q in questions:
                    lines.append(f"    • {q}")
    elif isinstance(protocol, str):
        lines.append(protocol)

    lines.append("")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# 3. generate_report_docx
# ---------------------------------------------------------------------------

async def generate_report_docx(report_json: dict) -> bytes:
    """Build a professional DOCX report from report_json and return as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    def _set_font(run, size: int = 11, bold: bool = False) -> None:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold

    def _heading(doc, text: str, level: int = 2) -> None:
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            _set_font(run, size=14 if level == 1 else 13, bold=True)

    def _para(doc, text: str, size: int = 11) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_font(run, size=size)

    axis_v = report_json.get("axis_vector", {})
    layer_v = report_json.get("layer_vector", {})
    dominant = report_json.get("dominant_cells", [])
    tension = report_json.get("tension_matrix", {})
    rigidity = report_json.get("rigidity", {})
    confidence = report_json.get("confidence", 0.0)
    structural = report_json.get("structural_report", "")
    protocol = report_json.get("interview_protocol", {})
    ts = report_json.get("timestamp", "")
    assessment_id = report_json.get("assessment_id", "—")
    try:
        date_str = datetime.fromisoformat(ts).strftime("%d.%m.%Y")
    except Exception:
        date_str = ts[:10] if ts else "—"

    # Title
    title_p = doc.add_heading("PsycheOS Screening v2 — Структурный профиль", level=1)
    for run in title_p.runs:
        _set_font(run, size=16, bold=True)

    _para(doc, f"Дата: {date_str}    ID: {assessment_id}")
    _para(doc, f"Уверенность: {confidence * 100:.0f}%")
    doc.add_paragraph()

    # Axis profile table
    _heading(doc, "Профиль осей регуляции")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Ось", "Значение", "Описание"]):
        run = cell.paragraphs[0].add_run(text)
        _set_font(run, bold=True)
    for akey in ["A1", "A2", "A3", "A4"]:
        val = axis_v.get(akey, 0.0)
        meta = _AXIS_META[akey]
        label = _axis_label(akey, val)
        row = table.add_row().cells
        _set_font(row[0].paragraphs[0].add_run(f"{meta['name']} ({akey})"))
        _set_font(row[1].paragraphs[0].add_run(f"{val:+.3f}"))
        _set_font(row[2].paragraphs[0].add_run(label))

    doc.add_paragraph()

    # Layer table
    _heading(doc, "Доминирующие слои")
    layer_table = doc.add_table(rows=1, cols=3)
    layer_table.style = "Table Grid"
    lhdr = layer_table.rows[0].cells
    for cell, text in zip(lhdr, ["Ранг", "Слой", "Значение"]):
        run = cell.paragraphs[0].add_run(text)
        _set_font(run, bold=True)
    sorted_layers = sorted(layer_v.items(), key=lambda kv: abs(kv[1]), reverse=True)
    for rank, (lkey, lval) in enumerate(sorted_layers, start=1):
        layer_name = _LAYER_META.get(lkey, lkey)
        row = layer_table.add_row().cells
        _set_font(row[0].paragraphs[0].add_run(str(rank)))
        _set_font(row[1].paragraphs[0].add_run(f"{layer_name} ({lkey})"))
        _set_font(row[2].paragraphs[0].add_run(f"{lval:+.3f}"))

    doc.add_paragraph()

    # Dominant cells
    _heading(doc, "Ключевые сочетания L×A")
    for rank, cell in enumerate(dominant[:3], start=1):
        cell_val = tension.get(cell, 0.0)
        _para(doc, f"{rank}. {cell}: {cell_val:+.3f}")

    doc.add_paragraph()

    # Rigidity
    _heading(doc, "Индекс гибкости")
    rig_total = rigidity.get("total", 0.0)
    rig_label = _rigidity_label(rig_total)
    _para(doc, f"Ригидность: {rig_label} ({rig_total:.2f})")
    _para(
        doc,
        f"Поляризация: {rigidity.get('polarization', 0.0):.2f}  "
        f"Стабильность: {rigidity.get('low_variance', 0.0):.2f}  "
        f"Повторяемость: {rigidity.get('strategy_repetition', 0.0):.2f}",
    )

    doc.add_paragraph()

    # Structural report
    _heading(doc, "Пояснение")
    _para(doc, structural)

    doc.add_paragraph()

    # Interview protocol
    _heading(doc, "Ориентиры для первой сессии")
    if isinstance(protocol, dict):
        for section_key, section_label in [
            ("axis_verification", "Верификация осей"),
            ("layer_exploration", "Исследование слоёв"),
            ("functional_context", "Функциональный контекст"),
        ]:
            questions = protocol.get(section_key, [])
            if questions:
                _para(doc, section_label + ":", size=12)
                for q in questions:
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(q)
                    _set_font(run)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

async def _call_claude(
    client,
    system: str,
    user_content: str,
    model: str = _SONNET,
    max_tokens: int = 2000,
) -> str | None:
    try:
        resp = await client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": user_content}],
            temperature=0.4,
            top_p=0.95,
        )
        return resp.content[0].text
    except Exception:
        logger.exception("[screen] Report Claude API call failed")
        return None


def _parse_json(text: str) -> dict:
    t = text.strip()
    if t.startswith("```json"):
        t = t[7:]
    if t.startswith("```"):
        t = t[3:]
    if t.endswith("```"):
        t = t[:-3]
    return json.loads(t.strip())
