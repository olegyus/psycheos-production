"""Генератор отчёта v1.1 — текст Claude + метрики → .docx (BytesIO).

v1.1:
  - TSI-компоненты на титульной странице
  - Таблица сигналов по репликам из iteration_log
  - Альтернативные траектории
  - Профиль специалиста (если есть история)
  - CCI кейса
"""

import io
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.simulator.schemas import (
    IterationLog, TSIComponents, CCIComponents,
    SpecialistProfile,
)


def generate_report_docx(
    report_text: str,
    case_name: str,
    case_id: str,
    session_goal: str,
    mode: str,
    crisis_flag: str,
    signal_log: list[str],
    fsm_log: list[str],
    iteration_log: list[IterationLog] = None,
    tsi: Optional[TSIComponents] = None,
    cci: Optional[CCIComponents] = None,
    specialist_profile: Optional[SpecialistProfile] = None,
) -> io.BytesIO:
    """Генерирует .docx отчёт. Возвращает BytesIO буфер."""

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(2)

    _add_title_page(doc, case_name, case_id, session_goal, mode,
                    crisis_flag, signal_log, fsm_log, tsi, cci)
    doc.add_page_break()

    if iteration_log:
        _add_signal_table(doc, iteration_log)
        doc.add_page_break()

    _parse_and_add_content(doc, report_text, skip_tables=True)

    if specialist_profile and specialist_profile.sessions_count > 1:
        doc.add_page_break()
        _add_specialist_profile(doc, specialist_profile)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

def _add_title_page(doc, case_name, case_id, session_goal, mode,
                    crisis_flag, signal_log, fsm_log, tsi, cci):

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("АНАЛИТИЧЕСКИЙ ОТЧЁТ")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PsycheOS Simulator v1.1")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    _add_separator(doc)
    doc.add_paragraph()

    meta = [
        ("Кейс", f"{case_name} ({case_id})"),
        ("Цель сессии", session_goal),
        ("Режим", mode),
        ("Кризисный флаг", crisis_flag),
        ("Дата", datetime.now().strftime("%d.%m.%Y %H:%M")),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.add_run(value).font.size = Pt(11)

    doc.add_paragraph()

    if tsi or cci:
        _add_separator(doc)
        doc.add_paragraph()

    if tsi:
        _add_tsi_block(doc, tsi)

    if cci:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CCI (Case Complexity Index): {cci.cci:.2f}")
        run.bold = True
        run.font.size = Pt(13)

        if tsi:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            comparison = "специалист справляется" if tsi.tsi >= cci.cci else "высокая вероятность дестабилизации"
            color = RGBColor(0x27, 0x8B, 0x37) if tsi.tsi >= cci.cci else RGBColor(0xCC, 0x33, 0x33)
            run = p.add_run(f"TSI vs CCI: {comparison}")
            run.font.size = Pt(11)
            run.font.color.rgb = color

    doc.add_paragraph()

    greens = signal_log.count("🟢")
    yellows = signal_log.count("🟡")
    reds = signal_log.count("🔴")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Сигналы: {greens} зелёных | {yellows} жёлтых | {reds} красных "
        f"(всего: {len(signal_log)})"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if fsm_log:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Траектория: {' → '.join(fsm_log[-15:])}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_tsi_block(doc, tsi: TSIComponents):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"TSI (Therapeutic Stability Index): {tsi.tsi:.2f}")
    run.bold = True
    run.font.size = Pt(18)

    if tsi.tsi >= 0.85:
        color = RGBColor(0x27, 0x8B, 0x37)
    elif tsi.tsi >= 0.70:
        color = RGBColor(0x33, 0x66, 0xCC)
    elif tsi.tsi >= 0.50:
        color = RGBColor(0xCC, 0x88, 0x00)
    else:
        color = RGBColor(0xCC, 0x33, 0x33)
    run.font.color.rgb = color

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Устойчивость позиции: {tsi.interpretation}")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = color

    doc.add_paragraph()

    components = [
        ("R_match (соответствие уровню)", tsi.R_match, 0.25),
        ("L_consistency (последовательность)", tsi.L_consistency, 0.20),
        ("Alliance_score (контакт)", tsi.Alliance_score, 0.20),
        ("Uncertainty_modulation (неопределённость)", tsi.Uncertainty_modulation, 0.20),
        ("Therapist_reactivity (реактивность, инв.)", tsi.Therapist_reactivity, 0.15),
    ]

    table = doc.add_table(rows=len(components) + 1, cols=3)
    table.style = "Light Grid Accent 1"

    for i, header in enumerate(["Параметр", "Значение", "Вес"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)

    for r_idx, (name, value, weight) in enumerate(components, 1):
        table.rows[r_idx].cells[0].paragraphs[0].add_run(name).font.size = Pt(10)

        val_run = table.rows[r_idx].cells[1].paragraphs[0].add_run(f"{value:.2f}")
        val_run.bold = True
        val_run.font.size = Pt(10)
        if value >= 0.80:
            val_run.font.color.rgb = RGBColor(0x27, 0x8B, 0x37)
        elif value >= 0.60:
            val_run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)
        else:
            val_run.font.color.rgb = RGBColor(0xCC, 0x88, 0x00)

        table.rows[r_idx].cells[2].paragraphs[0].add_run(f"{weight:.2f}").font.size = Pt(10)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# SIGNAL TABLE (from iteration_log)
# ═══════════════════════════════════════════════════════════════════════════

def _add_signal_table(doc, iteration_log: list[IterationLog]):
    _add_heading(doc, "Таблица сигналов по репликам", 2)

    headers = ["#", "FSM", "Layer", "Signal", "Match", "Cascade",
               "Δtrust", "ΔL0", "Δunc.", "Δdef.", "Причина"]
    num_cols = len(headers)

    table = doc.add_table(rows=len(iteration_log) + 1, cols=num_cols)
    table.style = "Light Grid Accent 1"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(8)

    signal_emoji = {"GREEN": "🟢", "YELLOW": "🟡", "RED": "🔴"}

    for r_idx, it in enumerate(iteration_log, 1):
        row = table.rows[r_idx]
        data = [
            str(it.replica_id),
            it.fsm_before,
            it.active_layer_before,
            signal_emoji.get(it.signal.value, "?"),
            f"{it.regulatory_match_score:.2f}",
            f"{it.cascade_probability:.2f}",
            f"{it.delta.trust:+d}",
            f"{it.delta.tension_L0:+d}",
            f"{it.delta.uncertainty:+d}",
            f"{it.delta.defense_activation:+d}",
            it.signal_reason[:40],
        ]
        for c_idx, val in enumerate(data):
            cell = row.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(8)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# SPECIALIST PROFILE
# ═══════════════════════════════════════════════════════════════════════════

def _add_specialist_profile(doc, profile: SpecialistProfile):
    _add_heading(doc, "Профиль специалиста (накопительная динамика)", 2)

    items = [
        ("Сессий проведено", str(profile.sessions_count)),
        ("Средний TSI", f"{profile.average_tsi:.2f}"),
        ("Средний Δtrust", f"{profile.average_delta_trust:+.1f}"),
        ("Доля жёлтых сигналов", f"{profile.yellow_ratio:.1%}"),
        ("Доля красных сигналов", f"{profile.red_ratio:.1%}"),
        ("Рекомендуемая сложность", f"CCI ≤ {profile.recommended_case_complexity:.2f}"),
    ]

    if profile.dominant_error_pattern:
        items.append(("Типичная ошибка", profile.dominant_error_pattern))
    if profile.typical_jump_level:
        items.append(("Типичный прыжок", profile.typical_jump_level))

    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Light Grid Accent 1"

    for r_idx, (label, value) in enumerate(items):
        row = table.rows[r_idx]
        run = row.cells[0].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(value).font.size = Pt(10)

    if len(profile.tsi_history) > 1:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Динамика TSI: ")
        run.bold = True
        values = " → ".join(f"{v:.2f}" for v in profile.tsi_history)
        p.add_run(values)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# CONTENT PARSER (markdown → docx)
# ═══════════════════════════════════════════════════════════════════════════

def _parse_and_add_content(doc, text, skip_tables: bool = False):
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            _add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue

        if stripped.startswith("## "):
            _add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue

        if stripped.startswith("### "):
            _add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue

        if re.match(r'^[═─\-]{5,}$', stripped):
            _add_separator(doc)
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            _add_code_block(doc, "\n".join(code_lines))
            continue

        if "|" in stripped and i + 1 < len(lines) and "|" in lines[i + 1]:
            table_lines = []
            while i < len(lines) and "|" in lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1
            if not skip_tables:
                _add_table(doc, table_lines)
            continue

        if stripped.startswith(("- ", "• ", "* ", "— ")):
            items = []
            while i < len(lines) and lines[i].strip().startswith(("- ", "• ", "* ", "— ")):
                item_text = re.sub(r'^[\-•\*—]\s*', '', lines[i].strip())
                items.append(item_text)
                i += 1
            _add_bullet_list(doc, items)
            continue

        _add_paragraph_with_formatting(doc, stripped)
        i += 1


def _add_heading(doc, text, level):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 12))


def _add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 50)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)


def _add_code_block(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _add_runs_with_bold(p, item)


def _add_paragraph_with_formatting(doc, text):
    p = doc.add_paragraph()
    _add_runs_with_bold(p, text)


def _add_runs_with_bold(paragraph, text):
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_table(doc, table_lines):
    data_lines = [
        line for line in table_lines
        if not re.match(r'^\|[\s\-:|\+]+\|$', line)
    ]
    if not data_lines:
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.split("|")]
        cells = [c for c in cells if c]
        rows.append(cells)

    if not rows:
        return

    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_runs_with_bold(p, cell_text)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True

    doc.add_paragraph()
